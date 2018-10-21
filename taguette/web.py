import asyncio
import hashlib
import hmac
import io
import os
from datetime import datetime
import json
import logging
import jinja2
import pkg_resources
from sqlalchemy.orm import aliased, joinedload, undefer, make_transient
from tornado.concurrent import Future
import tornado.ioloop
from tornado.routing import URLSpec
from tornado.web import authenticated, HTTPError, RequestHandler

from . import convert
from . import database
from .extract import extract


logger = logging.getLogger(__name__)


class BaseHandler(RequestHandler):
    """Base class for all request handlers.
    """
    template_env = jinja2.Environment(
        loader=jinja2.FileSystemLoader(
            [pkg_resources.resource_filename('taguette', 'templates')]
        ),
        autoescape=jinja2.select_autoescape(['html'])
    )

    @jinja2.contextfunction
    def _tpl_static_url(context, path):
        v = not context['handler'].application.settings.get('debug', False)
        return context['handler'].static_url(path, include_version=v)
    template_env.globals['static_url'] = _tpl_static_url

    @jinja2.contextfunction
    def _tpl_reverse_url(context, path, *args):
        return context['handler'].reverse_url(path, *args)
    template_env.globals['reverse_url'] = _tpl_reverse_url

    @jinja2.contextfunction
    def _tpl_xsrf_form_html(context):
        return jinja2.Markup(context['handler'].xsrf_form_html())
    template_env.globals['xsrf_form_html'] = _tpl_xsrf_form_html

    def __init__(self, application, request, **kwargs):
        super(BaseHandler, self).__init__(application, request, **kwargs)
        self.db = application.DBSession()

    def get_current_user(self):
        user = self.get_secure_cookie('user')
        if user is not None:
            return user.decode('utf-8')
        else:
            return None

    def login(self, username):
        logger.info("Logged in as %r", username)
        self.set_secure_cookie('user', username)

    def logout(self):
        logger.info("Logged out")
        self.clear_cookie('user')

    def render_string(self, template_name, **kwargs):
        template = self.template_env.get_template(template_name)
        return template.render(
            handler=self,
            current_user=self.current_user,
            multiuser=self.application.multiuser,
            register_enabled=self.application.register_enabled,
            **kwargs)

    def get_project(self, project_id):
        try:
            project_id = int(project_id)
        except ValueError:
            raise HTTPError(404)
        project_member = (
            self.db.query(database.ProjectMember)
            .options(joinedload(database.ProjectMember.project))
            .get((project_id, self.current_user))
        )
        if project_member is None:
            raise HTTPError(404)
        return project_member.project

    def get_document(self, project_id, document_id, contents=False):
        try:
            project_id = int(project_id)
            document_id = int(document_id)
        except ValueError:
            raise HTTPError(404)

        q = (
            self.db.query(database.Document)
            .options(joinedload(database.Document.project)
                     .joinedload(database.Project.members))
            .filter(database.Project.id == project_id)
            .filter(database.ProjectMember.user_login == self.current_user)
            .filter(database.Document.id == document_id)
        )
        if contents:
            q = q.options(undefer(database.Document.contents))
        document = q.one_or_none()
        if document is None:
            raise HTTPError(404)
        return document

    def get_json(self):
        type_ = self.request.headers.get('Content-Type', '')
        if not type_.startswith('application/json'):
            raise HTTPError(400, "Expected JSON")
        return json.loads(self.request.body.decode('utf-8'))

    def send_json(self, obj):
        if isinstance(obj, list):
            obj = {'results': obj}
        elif not isinstance(obj, dict):
            raise ValueError("Can't encode %r to JSON" % type(obj))
        self.set_header('Content-Type', 'application/json; charset=utf-8')
        return self.finish(json.dumps(obj))


class Index(BaseHandler):
    """Index page, shows welcome message and user's projects.
    """
    def get(self):
        if self.current_user is not None:
            user = self.db.query(database.User).get(self.current_user)
            if user is None:
                logger.warning("User is logged in as non-existent user %r",
                               self.current_user)
                self.logout()
            else:
                self.render('index.html', user=user, projects=user.projects)
                return
        elif not self.application.multiuser:
            token = self.get_query_argument('token', None)
            if token and token == self.application.single_user_token:
                self.login('admin')
                self.redirect(self.reverse_url('index'))
            else:
                self.render('token_needed.html')
        else:
            self.render('welcome.html')


class Login(BaseHandler):
    def get(self):
        if not self.application.multiuser:
            raise HTTPError(404)
        if self.current_user:
            self._go_to_next()
        else:
            self.render('login.html', register=False,
                        next=self.get_argument('next', ''))

    def post(self):
        login = self.get_body_argument('login')
        password = self.get_body_argument('password')
        user = self.db.query(database.User).get(login)
        if user is not None and user.check_password(password):
            self.login(user.login)
            self._go_to_next()
        else:
            self.render('login.html', register=False,
                        next=self.get_argument('next', ''),
                        login_error="Invalid login or password")

    def _go_to_next(self):
        next_ = self.get_argument('next')
        if not next_:
            next_ = self.reverse_url('index')
        self.redirect(next_)


class Logout(BaseHandler):
    def get(self):
        if not self.application.multiuser:
            raise HTTPError(404)
        self.logout()
        self.redirect(self.reverse_url('index'))


class Register(BaseHandler):
    def get(self):
        if not self.application.multiuser:
            raise HTTPError(404)
        if not self.application.register_enabled:
            raise HTTPError(403)
        if self.current_user:
            self.redirect(self.reverse_url('index'))
        else:
            self.render('login.html', register=True)

    def post(self):
        if not self.application.register_enabled:
            raise HTTPError(403)
        login = self.get_body_argument('login')
        password1 = self.get_body_argument('password1')
        password2 = self.get_body_argument('password2')
        if password1 != password2:
            self.render('login.html', register=True,
                        register_error="Passwords do not match")
            return
        if self.db.query(database.User).get(login) is not None:
            self.render('login.html', register=True,
                        register_error="Username is taken")
            return
        user = database.User(login=login)
        user.set_password(password1)
        self.db.add(user)
        self.db.commit()
        logger.info("User registered: %r", login)
        self.set_secure_cookie('user', login)
        self.redirect(self.reverse_url('index'))


class NewProject(BaseHandler):
    @authenticated
    def get(self):
        self.render('project_new.html')

    @authenticated
    def post(self):
        name = self.get_body_argument('name', '')
        description = self.get_body_argument('description', '')
        if not name:
            return self.render('project_new.html',
                               name=name, description=description,
                               error="Please enter a name")

        # Create project
        project = database.Project(name=name, description=description)
        self.db.add(project)
        # Add user as admin
        membership = database.ProjectMember(
            project=project,
            user_login=self.current_user,
            privileges=database.Privileges.ADMIN
        )
        self.db.add(membership)
        # Add default set of tags
        self.db.add(database.Tag(project=project, path='interesting',
                                 description="Further review required"))
        self.db.add(database.Tag(project=project, path='people',
                                 description="Known people"))

        self.db.commit()
        self.redirect(self.reverse_url('project', project.id))

    def render(self, template_name, **kwargs):
        for name in ('name', 'description', 'error'):
            kwargs.setdefault(name, '')
        super(NewProject, self).render(template_name, **kwargs)


class Project(BaseHandler):
    @authenticated
    def get(self, project_id):
        project = self.get_project(project_id)
        documents_json = jinja2.Markup(json.dumps(
            {
                str(doc.id): {'id': doc.id, 'name': doc.name,
                              'description': doc.description}
                for doc in project.documents
            },
            sort_keys=True,
        ))
        tags_json = jinja2.Markup(json.dumps(
            {
                str(tag.id): {'id': tag.id,
                              'path': tag.path,
                              'description': tag.description}
                for tag in project.tags
            },
            sort_keys=True,
        ))
        self.render('project.html',
                    project=project,
                    last_event=js_timestamp(project.last_event),
                    documents=documents_json,
                    tags=tags_json)


def js_timestamp(dt=None):
    if dt is None:
        dt = datetime.utcnow()
    return int(dt.timestamp())


class ProjectMeta(BaseHandler):
    @authenticated
    def post(self, project_id):
        obj = self.get_json()
        project = self.get_project(project_id)
        project.name = obj['name']
        project.description = obj['description']
        now = datetime.utcnow()
        project.meta_updated = now
        logger.info("Updated project: %r %r",
                    project.name, project.description)
        cmd = database.Command.project_meta(
            self.current_user,
            project.id,
            obj['name'],
            obj['description'],
        )
        self.db.add(cmd)
        self.db.commit()
        self.db.refresh(cmd)
        self.application.notify_project(project.id, cmd)
        return self.send_json({'ts': js_timestamp(now)})


class DocumentAdd(BaseHandler):
    @authenticated
    async def post(self, project_id):
        project = self.get_project(project_id)

        name = self.get_body_argument('name')
        description = self.get_body_argument('description')
        file = self.request.files['file'][0]
        content_type = file.content_type

        try:
            body = await convert.to_html_chunks(file.body, content_type,
                                                file.filename)
        except convert.ConversionError as err:
            self.set_status(400)
            self.send_json({
                'error': str(err),
            })
        else:
            doc = database.Document(
                name=name,
                description=description,
                filename=file.filename or None,
                project=project,
                contents=body,
            )
            self.db.add(doc)
            project.documents_updated = datetime.utcnow()
            cmd = database.Command.document_add(
                self.current_user,
                doc,
            )
            self.db.add(cmd)
            logger.info("Document added to project %r: %r %r (%d bytes)",
                        project.id, doc.id, doc.name, len(doc.contents))
            self.db.commit()
            self.db.refresh(cmd)
            self.application.notify_project(project.id, cmd)
            self.send_json({'created': doc.id})


class DocumentContents(BaseHandler):
    @authenticated
    def get(self, project_id, document_id):
        document = self.get_document(project_id, document_id, True)
        self.send_json({
            'contents': [
                {'offset': 0, 'contents': document.contents},
            ],
            'highlights': [
                {'id': hl.id,
                 'start_offset': hl.start_offset,
                 'end_offset': hl.end_offset,
                 'tags': [t.id for t in hl.tags]}
                for hl in document.highlights
            ],
        })


class HighlightAdd(BaseHandler):
    @authenticated
    def post(self, project_id, document_id):
        obj = self.get_json()
        document = self.get_document(project_id, document_id, True)
        start, end = obj['start_offset'], obj['end_offset']
        snippet = extract(document.contents, start, end)
        hl = database.Highlight(document=document,
                                start_offset=start,
                                end_offset=end,
                                snippet=snippet)
        self.db.add(hl)
        self.db.commit()  # Need to commit to get hl.id
        self.db.bulk_insert_mappings(database.HighlightTag, [
            dict(
                highlight_id=hl.id,
                tag_id=tag,
            )
            for tag in obj.get('tags', [])
        ])
        cmd = database.Command.highlight_add(
            self.current_user,
            document,
            hl,
            obj.get('tags', []),
        )
        self.db.add(cmd)
        self.db.commit()
        self.db.refresh(cmd)
        self.application.notify_project(document.project_id, cmd)

        self.send_json({'id': hl.id})


class HighlightUpdate(BaseHandler):
    @authenticated
    def post(self, project_id, document_id, highlight_id):
        obj = self.get_json()
        document = self.get_document(project_id, document_id)
        hl = self.db.query(database.Highlight).get(int(highlight_id))
        if hl.document_id != document.id:
            raise HTTPError(404)
        if not obj:
            self.db.delete(hl)
            cmd = database.Command.highlight_delete(
                self.current_user,
                document,
                hl.id,
            )
        else:
            if 'start_offset' in obj:
                hl.start_offset = obj['start_offset']
            if 'end_offset' in obj:
                hl.end_offset = obj['end_offset']
            if 'tags' in obj:
                (
                    self.db.query(database.HighlightTag)
                    .filter(database.HighlightTag.highlight == hl)
                ).delete()
                self.db.bulk_insert_mappings(database.HighlightTag, [
                    dict(
                        highlight_id=hl.id,
                        tag_id=tag,
                    )
                    for tag in obj.get('tags', [])
                ])
            cmd = database.Command.highlight_add(
                self.current_user,
                document,
                hl,
                obj.get('tags', []),
            )
        self.db.add(cmd)
        self.db.commit()
        self.db.refresh(cmd)
        self.application.notify_project(document.project_id, cmd)

        self.send_json({'id': hl.id})


class Highlights(BaseHandler):
    @authenticated
    def get(self, project_id, path):
        project = self.get_project(project_id)

        tag = aliased(database.Tag)
        hltag = aliased(database.HighlightTag)
        highlights = (
            self.db.query(database.Highlight)
            .join(hltag, hltag.highlight_id == database.Highlight.id)
            .join(tag, hltag.tag_id == tag.id)
            .filter(tag.path.startswith(path))
            .filter(tag.project == project)
        ).all()

        self.send_json({
            'highlights': [
                {
                    'id': hl.id,
                    'document_id': hl.document_id,
                    'content': hl.snippet,
                }
                for hl in highlights
            ],
        })


class TagAdd(BaseHandler):
    @authenticated
    def post(self, project_id):
        obj = self.get_json()
        project = self.get_project(project_id)
        tag = database.Tag(project=project,
                           path=obj['path'], description=obj['description'])
        self.db.add(tag)
        self.db.commit()  # Need to commit to get tag.id
        cmd = database.Command.tag_add(
            self.current_user,
            tag,
        )
        self.db.add(cmd)
        self.db.commit()
        self.db.refresh(cmd)
        self.application.notify_project(project.id, cmd)

        self.send_json({'id': tag.id})


class TagUpdate(BaseHandler):
    @authenticated
    def post(self, project_id, tag_id):
        obj = self.get_json()
        project = self.get_project(project_id)
        tag = self.db.query(database.Tag).get(int(tag_id))
        if tag.project_id != project.id:
            raise HTTPError(404)
        if not obj:
            self.db.delete(tag)
            cmd = database.Command.tag_delete(
                self.current_user,
                project.id,
                tag.id,
            )
        else:
            if 'path' in obj:
                tag.path = obj['path']
            if 'description' in obj:
                tag.description = obj['description']
            cmd = database.Command.tag_add(
                self.current_user,
                tag,
            )
        self.db.add(cmd)
        self.db.commit()
        self.db.refresh(cmd)
        self.application.notify_project(project.id, cmd)

        self.send_json({'id': tag.id})


class ProjectEvents(BaseHandler):
    @authenticated
    async def get(self, project_id):
        from_ts = int(self.get_query_argument('from'))
        from_dt = datetime.utcfromtimestamp(from_ts)
        project = self.get_project(project_id)
        self.project_id = int(project_id)

        # Check for immediate update
        cmd = (
            self.db.query(database.Command)
            .filter(database.Command.date > from_dt)
        ).one_or_none()

        # Wait for an event
        if cmd is None:
            self.wait_future = Future()
            self.application.observe_project(project.id, self.wait_future)
            self.db.expire_all()
            try:
                cmd = await self.wait_future
            except asyncio.CancelledError:
                return

        payload = dict(cmd.payload)
        type_ = payload.pop('type', None)
        if type_ == 'project_meta':
            result = {'project_meta': payload}
        elif type_ == 'document_add':
            payload['id'] = cmd.document_id
            result = {'document_add': [payload]}
        elif type_ == 'highlight_add':
            result = {'highlight_add': {cmd.document_id: [payload]}}
        elif type_ == 'highlight_delete':
            result = {
                'highlight_delete': {cmd.document_id: [payload['id']]}
            }
        elif type_ == 'tag_add':
            result = {
                'tag_add': [payload],
            }
        elif type_ == 'tag_delete':
            result = {
                'tag_delete': [payload['id']],
            }
        else:
            raise ValueError("Unknown command type %r" % type_)

        result['ts'] = js_timestamp(cmd.date)
        self.send_json(result)

    def on_connection_close(self):
        self.wait_future.cancel()
        self.application.unobserve_project(self.project_id, self.wait_future)


class ExportCodebookCsv(BaseHandler):
    def get(self, project_id):
        import csv

        project = self.get_project(project_id)
        tags = list(project.tags)
        self.set_header('Content-Type', 'text/csv; charset=utf-8')
        self.set_header('Content-Disposition',
                        'attachment; filename="codebook.csv"')
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(['tag', 'description'])
        for tag in tags:
            writer.writerow([tag.path, tag.description])
        self.finish(buf.getvalue())


class ExportCodebookDoc(BaseHandler):
    def get(self, project_id):
        from docx import Document

        project = self.get_project(project_id)
        tags = list(project.tags)
        self.set_header(
            'Content-Type',
            'application/vnd.openxmlformats-officedocument.wordprocessingml'
            '.document; charset=utf-8')
        self.set_header('Content-Disposition',
                        'attachment; filename="codebook.docx"')
        buf = io.BytesIO()
        document = Document()
        document.add_heading('Taguette Codebook', 0)
        for tag in tags:
            document.add_heading(tag.path, 1)
            document.add_paragraph(tag.description)
        document.save(buf)
        self.finish(buf.getvalue())


class Application(tornado.web.Application):
    def __init__(self, handlers, db_url, multiuser, cookie_secret,
                 register_enabled=True, **kwargs):
        # Don't reuse the secret
        cookie_secret = cookie_secret + (
            '.multi' if multiuser
            else '.single'
        )

        super(Application, self).__init__(handlers,
                                          cookie_secret=cookie_secret,
                                          **kwargs)

        self.multiuser = multiuser
        self.register_enabled = register_enabled

        self.DBSession = database.connect(db_url)
        self.event_waiters = {}

        db = self.DBSession()
        admin = (
            db.query(database.User)
            .filter(database.User.login == 'admin')
            .one_or_none()
        )
        if admin is None:
            logger.warning("Creating user 'admin'")
            admin = database.User(login='admin')
            if self.multiuser:
                self._set_password(admin)
            db.add(admin)
            db.commit()
        elif self.multiuser and not admin.hashed_password:
            self._set_password(admin)
            db.commit()

        if self.multiuser:
            self.single_user_token = None
            logging.info("Starting in multi-user mode")
        else:
            self.single_user_token = hmac.new(
                cookie_secret.encode('utf-8'),
                b'taguette_single_user',
                digestmod=hashlib.sha256,
            ).hexdigest()

    def _set_password(self, user):
        import getpass
        passwd = getpass.getpass("Enter password for user %r: " % user.login)
        user.set_password(passwd)

    def observe_project(self, project_id, future):
        assert isinstance(project_id, int)
        self.event_waiters.setdefault(project_id, set()).add(future)

    def unobserve_project(self, project_id, future):
        assert isinstance(project_id, int)
        self.event_waiters[project_id].remove(future)

    def notify_project(self, project_id, cmd):
        assert isinstance(project_id, int)
        make_transient(cmd)
        for future in self.event_waiters.pop(project_id, []):
            future.set_result(cmd)


def make_app(db_url, multiuser, register_enabled=True, debug=False):
    if 'XDG_CACHE_HOME' in os.environ:
        cache = os.environ['XDG_CACHE_HOME']
    else:
        cache = os.path.expanduser('~/.cache')
    os.makedirs(cache, 0o700, exist_ok=True)
    cache = os.path.join(cache, 'taguette.json')
    secret = None
    try:
        fp = open(cache)
    except IOError:
        pass
    else:
        try:
            secret = json.load(fp)['cookie_secret']
            fp.close()
        except Exception:
            logger.exception("Couldn't load cookie secret from cache file")
        if not isinstance(secret, str) or not 10 <= len(secret) < 2048:
            logger.error("Invalid cookie secret in cache file")
            secret = None
    if secret is None:
        secret = os.urandom(30).decode('iso-8859-15')
        try:
            fp = open(cache, 'w')
            json.dump({'cookie_secret': secret}, fp)
            fp.close()
        except IOError:
            logger.error("Couldn't open cache file, cookie secret won't be "
                         "persisted! Users will be logged out if you restart "
                         "the program.")

    return Application(
        [
            URLSpec('/', Index, name='index'),
            URLSpec('/login', Login, name='login'),
            URLSpec('/logout', Logout, name='logout'),
            URLSpec('/register', Register, name='register'),
            URLSpec('/new', NewProject, name='new_project'),
            URLSpec('/project/([0-9]+)', Project, name='project'),
            URLSpec('/project/([0-9]+)/document/[0-9]+', Project,
                    name='project_doc'),
            URLSpec('/project/([0-9]+)/highlights/[^/]+', Project,
                    name='project_tag'),
            URLSpec('/project/([0-9]+)/meta', ProjectMeta),
            URLSpec('/project/([0-9]+)/document/new', DocumentAdd),
            URLSpec('/project/([0-9]+)/document/([0-9]+)/content',
                    DocumentContents),
            URLSpec('/project/([0-9]+)/document/([0-9]+)/highlight/new',
                    HighlightAdd),
            URLSpec('/project/([0-9]+)/document/([0-9]+)/highlight/([0-9]+)',
                    HighlightUpdate),
            URLSpec('/project/([0-9]+)/highlights/([^/]+)/list', Highlights),
            URLSpec('/project/([0-9]+)/tag/new', TagAdd),
            URLSpec('/project/([0-9]+)/tag/([0-9]+)', TagUpdate),
            URLSpec('/project/([0-9]+)/events', ProjectEvents),
            URLSpec('/project/([0-9]+)/export/codebook.csv', ExportCodebookCsv,
                    name='export_codebook_csv'),
            URLSpec('/project/([0-9]+)/export/codebook.docx',
                    ExportCodebookDoc, name='export_codebook_doc'),
        ],
        static_path=pkg_resources.resource_filename('taguette', 'static'),
        login_url='/login',
        xsrf_cookies=True,
        debug=debug,
        multiuser=multiuser,
        register_enabled=register_enabled,
        cookie_secret=secret,
        db_url=db_url,
    )
